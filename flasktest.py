# -*- coding: utf-8 -*-
"""
Created on Thu Feb 13 20:39:15 2025

@author: Mark
"""

import os
import uuid
from datetime import datetime
from flask import Flask, render_template, redirect, url_for, flash, session, request
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_sqlalchemy import SQLAlchemy
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import InputRequired, Length, Email, ValidationError
from flask_bcrypt import Bcrypt
from flask_wtf.file import FileField, FileAllowed
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from docx import Document as DocxDocument
import PyPDF2

# Load environment variables
load_dotenv()

app = Flask(__name__)

DB_USER = 'root'
DB_PASSWORD = 'Arsenal%401957'  # Corrected password (URL-encoded '@' as '%40')
DB_HOST = 'localhost'
DB_PORT = '3306'
DB_NAME = 'document_storage'
UPLOAD_FOLDER = os.getenv('UPLOAD_FOLDER', 'uploads/')

app.config['SECRET_KEY'] = 'thisisasecretkey'
app.config['SQLALCHEMY_DATABASE_URI'] = f'mysql+pymysql://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
db = SQLAlchemy(app)
bcrypt = Bcrypt(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

class User(db.Model, UserMixin): 
    user_id = db.Column(db.Integer, primary_key=True, autoincrement=True) 
    username = db.Column(db.String(100), unique=True, nullable=False) 
    email = db.Column(db.String(255), unique=True, nullable=False) 
    password = db.Column(db.String(255), nullable=False) 
    role = db.Column(db.Enum('admin', 'user'), default='user', nullable=False) 

    def get_id(self):
        return str(self.user_id)  # Override to use user_id instead of id
    
class Document(db.Model):
    doc_id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    title = db.Column(db.String(255), nullable=False)
    author = db.Column(db.String(255), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    file_path = db.Column(db.String(255), nullable=False)
    file_type = db.Column(db.String(50), nullable=False)
    size_kb = db.Column(db.Integer, nullable=False)
    uploaded_by = db.Column(db.Integer, db.ForeignKey('user.user_id'), nullable=False)
    user = db.relationship('User', backref=db.backref('document', lazy=True))
    content = db.relationship('Content', backref='document', uselist=False, lazy=True)

class Content(db.Model):
    content_id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    doc_id = db.Column(db.Integer, db.ForeignKey('document.doc_id'), nullable=False)
    text_content = db.Column(db.Text, nullable=True)

class SearchLogs(db.Model):
    search_id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    query = db.Column(db.String(255), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.user_id'), nullable=False)
    search_date = db.Column(db.DateTime, default=datetime.utcnow)


class RegisterForm(FlaskForm):
    username = StringField(validators=[InputRequired(), Length(min=4, max=20)], render_kw={"placeholder": "Username"})
    email = StringField(validators=[InputRequired(), Email(), Length(max=255)], render_kw={"placeholder": "Email"})
    password = PasswordField(validators=[InputRequired(), Length(min=8, max=20)], render_kw={"placeholder": "Password"})
    submit = SubmitField('Register')

    def validate_username(self, username):
        existing_user = User.query.filter_by(username=username.data).first()
        if existing_user:
            raise ValidationError('That username already exists. Please choose a different one.')

class LoginForm(FlaskForm):
    username = StringField(validators=[InputRequired(), Length(min=4, max=20)], render_kw={"placeholder": "Username"})
    password = PasswordField(validators=[InputRequired(), Length(min=8, max=20)], render_kw={"placeholder": "Password"})
    submit = SubmitField('Login')

class UploadForm(FlaskForm):
    title = StringField('Document Title', validators=[InputRequired(), Length(max=255)])
    author = StringField('Author', validators=[InputRequired(), Length(max=255)])
    file = FileField('Upload Document', validators=[InputRequired(), FileAllowed(['pdf', 'docx', 'txt'], 'Documents only!')])
    submit = SubmitField('Upload')

class SearchForm(FlaskForm):
    query = StringField('Search', validators=[InputRequired(), Length(min=1, max=255)], render_kw={"placeholder": "Search..."})
    submit = SubmitField('Search')

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('upload'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and bcrypt.check_password_hash(user.password, form.password.data):
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('upload'))
        else:
            flash('Invalid credentials.', 'danger')
    
    return render_template('login.html', form=form)

@app.route('/dashboard')
@login_required
def dashboard():
    return f"Welcome, {current_user.username}! <a href='{url_for('logout')}'>Logout</a>"

@app.route('/logout')
@login_required
def logout():
    logout_user()
    session.clear()
    flash("Logged out successfully.", "info")
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    form = RegisterForm()
    if form.validate_on_submit():
        hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
        new_user = User(username=form.username.data, email=form.email.data, password=hashed_password, role='user')
        db.session.add(new_user)
        db.session.commit()
        flash("Registration successful! Please log in.", "success")
        return redirect(url_for('login'))
    
    return render_template('register.html', form=form)


@app.route('/search', methods=['GET', 'POST'])
@login_required
def search():
    form = SearchForm()
    results = []

    if form.validate_on_submit():
        query = form.query.data.strip()
        print(f"🔍 Received search query: {query}")

        if not query:
            flash("Please enter a search term.", "warning")
            return render_template('search.html', form=form, results=[])

        # Log search
        log = SearchLogs(query=query, user_id=current_user.user_id)
        db.session.add(log)
        db.session.commit()

        # Perform search
        results = (
            db.session.query(Document)
            .join(Content)
            .filter(Content.text_content.ilike(f"%{query}%"))
            .order_by(Document.upload_date.desc())
            .all()
        )

        print(f"✅ Found {len(results)} results for '{query}'")

        if not results:
            flash("No results found.", "info")

    return render_template('search.html', form=form, results=results)


def read_file_contents(file_path, file_ext):
    if file_ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            print(f"Read TXT content: {content[:100]}...")  # Debugging output (first 100 chars)
            return content
    elif file_ext == 'pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            print(f"Read PDF content: {content[:100]}...")  # Debugging output (first 100 chars)
            return content
    elif file_ext == 'docx':
        doc = DocxDocument(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        print(f"Read DOCX content: {content[:100]}...")  # Debugging output (first 100 chars)
        return content
    return ""

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    form = UploadForm()
    if form.validate_on_submit():
        file = form.file.data
        filename = secure_filename(file.filename)

        # Ensure uploads folder is next to templates directory
        template_dir = os.path.join(app.root_path, 'templates')
        upload_folder = os.path.join(os.path.dirname(template_dir), 'uploads')

        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)

        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)

        file_ext = os.path.splitext(filename)[1][1:]  # Remove the dot
        file_size_kb = os.path.getsize(filepath) // 1024  # Convert bytes to KB

        new_document = Document(
            title=form.title.data,
            author=form.author.data,
            file_path=filepath,
            file_type=file_ext,
            size_kb=file_size_kb,
            uploaded_by=current_user.user_id  # Updated reference here
        )

        db.session.add(new_document)
        db.session.commit()

        # Now store content separately
        content_text = read_file_contents(filepath, file_ext)
        content_entry = Content(doc_id=new_document.doc_id, text_content=content_text)
        db.session.add(content_entry)
        db.session.commit()

        flash("File uploaded successfully!", "success")
        return redirect(url_for('upload'))
    return render_template('upload.html', form=form)

if __name__ == "__main__":
    app.run(debug=True)
