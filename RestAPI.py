from flask import Flask, request, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_bcrypt import Bcrypt
from flask_jwt_extended import (
    JWTManager, create_access_token, jwt_required
)
from datetime import datetime
import os
from docx import Document as DocxDocument
import PyPDF2

# Init app
app = Flask(__name__)
app.config['SECRET_KEY'] = 'thisisasecretkey'
app.config['JWT_SECRET_KEY'] = 'jwtsecretkey'
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:Arsenal%401957@localhost:3306/document_storage'
app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', 'uploads/')
db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
jwt = JWTManager(app)

# Models
class User(db.Model):
    user_id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(255), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    role = db.Column(db.Enum('admin', 'user'), default='user', nullable=False)

class Document(db.Model):
    doc_id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    author = db.Column(db.String(255), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    file_path = db.Column(db.String(255), nullable=False)
    file_type = db.Column(db.String(50), nullable=False)
    size_kb = db.Column(db.Integer, nullable=False)
    uploaded_by = db.Column(db.Integer, db.ForeignKey('user.user_id'), nullable=False)
    user = db.relationship('User', backref=db.backref('documents', lazy=True))

class Content(db.Model):
    content_id = db.Column(db.Integer, primary_key=True)
    doc_id = db.Column(db.Integer, db.ForeignKey('document.doc_id'), nullable=False)
    text_content = db.Column(db.Text)

# Utilities
def read_file_contents(file_path, file_ext):
    if file_ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_ext == 'pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_ext == 'docx':
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""

# Routes

@app.route("/api/register", methods=["POST"])
def register():
    data = request.get_json()
    if User.query.filter_by(username=data['username']).first():
        return jsonify({"msg": "Username already exists"}), 409
    hashed_pw = bcrypt.generate_password_hash(data['password']).decode('utf-8')
    new_user = User(username=data['username'], email=data['email'], password=hashed_pw)
    db.session.add(new_user)
    db.session.commit()
    return jsonify({"msg": "Registration successful"}), 201

@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json()
    user = User.query.filter_by(username=data['username']).first()
    if user and bcrypt.check_password_hash(user.password, data['password']):
        access_token = create_access_token(identity=user.user_id)
        return jsonify(access_token=access_token)
    return jsonify({"msg": "Invalid credentials"}), 401

@app.route("/api/documents", methods=["GET"])
@jwt_required()
def list_documents():
    docs = Document.query.all()
    return jsonify([{
        "doc_id": d.doc_id,
        "title": d.title,
        "author": d.author,
        "upload_date": d.upload_date.isoformat(),
        "file_type": d.file_type,
        "size_kb": d.size_kb
    } for d in docs])

@app.route("/api/documents/<int:doc_id>/content", methods=["GET"])
@jwt_required()
def get_document_content(doc_id):
    content = Content.query.filter_by(doc_id=doc_id).first()
    if not content:
        return jsonify({"msg": "Content not found"}), 404
    return jsonify({
        "doc_id": doc_id,
        "text_content": content.text_content[:1000]  # trimmed for preview
    })

@app.route("/api/search", methods=["GET"])
@jwt_required()
def search_documents():
    query = request.args.get("q", "")
    if not query:
        return jsonify({"msg": "Missing query"}), 400
    results = (
        db.session.query(Document)
        .join(Content)
        .filter(Content.text_content.ilike(f"%{query}%"))
        .all()
    )
    return jsonify([{
        "doc_id": d.doc_id,
        "title": d.title,
        "author": d.author,
        "upload_date": d.upload_date.isoformat()
    } for d in results])

# Run
if __name__ == "__main__":
    app.run(debug=True)
